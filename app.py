import os
import re
from io import BytesIO
from zipfile import ZipFile

import streamlit as st
from docx import Document
import docx
import PyPDF2
from docx.enum.style import WD_STYLE_TYPE

# ── Page setup ───────────────────────────────────────────────────────────────
st.set_page_config(page_title="Neogen Job Advert Generator", page_icon="🧪", layout="centered")
st.image("assets/neogen-logo-green.webp", use_container_width=True)
st.title("📄 Neogen Job Advert Generator")

# Prefer Streamlit Secrets; fall back to env var (don’t build client yet)
try:
    OPENAI_API_KEY = st.secrets["OPENAI_API_KEY"]
except Exception:
    OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")

st.caption("🔐 OpenAI API key detected." if OPENAI_API_KEY else "❌ No OpenAI API key found. Add it in Manage app → Settings → Secrets.")

# ── Helpers: extract text ────────────────────────────────────────────────────
def extract_text_from_docx(file) -> str:
    d = docx.Document(file)
    return "\n".join(p.text for p in d.paragraphs if p.text.strip())

def extract_text_from_pdf(file) -> str:
    reader = PyPDF2.PdfReader(file)
    parts = []
    for page in reader.pages:
        txt = page.extract_text() or ""
        if txt.strip():
            parts.append(txt)
    return "\n".join(parts).strip()

# ── Helpers: sanitise output (remove Markdown) ───────────────────────────────
def strip_markdown(s: str) -> str:
    s = re.sub(r"\*\*(.*?)\*\*", r"\1", s)
    s = re.sub(r"\*(.*?)\*", r"\1", s)
    s = re.sub(r"`(.*?)`", r"\1", s)
    s = re.sub(r"^\s{0,3}#{1,6}\s*", "", s, flags=re.MULTILINE)  # remove markdown headings
    s = s.replace("•", "-").replace("–", "-").replace("—", "-")
    s = re.sub(r"\r\n", "\n", s)
    return s.strip()

# ── House style prompt (explicit headers + bullets) ──────────────────────────
HOUSE_STYLE = """
Rewrite the job description as a polished Neogen job advert in a clear, professional, concise tone.

Output rules (very important):
- PLAIN TEXT ONLY (no Markdown markup).
- Use these exact section headers when you have content for them (spelled as below):
  1) Location:
  2) Essential Duties and Responsibilities:
  3) Education and Experience:
- The opening paragraph comes before “Location:”.
- Bullet points MUST start with a single hyphen and a space: "- ".
- Keep bullets crisp, action-led, present tense, one line where possible.
- Do not invent details, benefits or salary.
- End with this exact closing line on its own line:
Please press Apply to submit your application.
"""

# ── OpenAI generator (client created lazily) ─────────────────────────────────
def generate_neogen_advert(job_description: str) -> str:
    from openai import OpenAI  # lazy import so UI always renders
    client = OpenAI(api_key=OPENAI_API_KEY)

    prompt = f"""You are a professional HR copywriter at Neogen.

Using the HOUSE STYLE, rewrite the JOB DESCRIPTION into a Neogen-style job advert.
Do not add facts that aren't in the input. Keep it concise and scannable.
IMPORTANT: Output must be plain text only (no Markdown). Use "- " bullets and the exact section headers.

HOUSE STYLE:
\"\"\"{HOUSE_STYLE}\"\"\"

JOB DESCRIPTION:
\"\"\"{job_description}\"\"\""""

    resp = client.chat.completions.create(
        model="gpt-4o-mini",   # switch to "gpt-4o" for highest quality if desired
        messages=[
            {"role": "system", "content": "You are a precise HR copywriter who follows style guides faithfully."},
            {"role": "user",   "content": prompt},
        ],
        max_tokens=1200,
        temperature=0.35,
    )
    return resp.choices[0].message.content

# ── DOCX builder: bold headers + robust bullets ──────────────────────────────
# Accept both with/without colon, case-insensitive
SECTION_HEADERS = {
    "Location:",
    "Essential Duties and Responsibilities:",
    "Education and Experience:",
    "Education and Experience",
}

def to_docx_structured(text: str) -> BytesIO:
    """
    Build a DOCX with Neogen-style formatting:
    - Known section headers are bold.
    - Lines beginning with '- ' (or '• ') become bullets.
    - Fallback to a visible '• ' if the 'List Bullet' style isn't available.
    """
    doc = Document()

    bullet_style_name = "List Bullet"
    bullet_style_exists = bullet_style_name in [
        s.name for s in doc.styles if s.type == WD_STYLE_TYPE.PARAGRAPH
    ]

    def add_header(line: str) -> bool:
        norm = line.strip()
        for hdr in SECTION_HEADERS:
            if norm.lower().startswith(hdr.lower()):
                p = doc.add_paragraph()
                # write bold header token (ensure trailing colon if needed)
                hdr_token = hdr
                run = p.add_run(hdr_token)
                run.bold = True
                # tail after header (e.g., "Location: Hybrid – EMEAI")
                tail = norm[len(hdr):].lstrip()
                if tail.startswith(":"):
                    tail = tail[1:].lstrip()
                if tail:
                    p.add_run(" " + tail)
                return True
        return False

    lines = [ln.rstrip() for ln in text.split("\n")]

    for raw in lines:
        line = raw.strip()

        if not line:
            doc.add_paragraph()
            continue

        if add_header(line):
            continue

        if line.startswith("- ") or line.startswith("• "):
            bullet_text = line[2:].strip()
            if bullet_style_exists:
                try:
                    doc.add_paragraph(bullet_text, style=bullet_style_name)
                except KeyError:
                    # very rare; fallback to glyph
                    p = doc.add_paragraph()
                    p.add_run("• ").bold = True
                    p.add_run(bullet_text)
            else:
                # glyph fallback
                p = doc.add_paragraph()
                p.add_run("• ").bold = True
                p.add_run(bullet_text)
            continue

        doc.add_paragraph(line)

    out = BytesIO()
    doc.save(out)
    out.seek(0)
    return out

# ── UI: tabs (single & batch) ────────────────────────────────────────────────
tab_single, tab_batch = st.tabs(["Single file", "Batch (multiple files)"])

with tab_single:
    uploaded = st.file_uploader(
        "Upload a Job Description (.docx or .pdf)",
        type=["docx", "pdf"],
        key="single_uploader"
    )

    if uploaded:
        try:
            jd_text = extract_text_from_docx(uploaded) if uploaded.name.lower().endswith(".docx") \
                      else extract_text_from_pdf(uploaded)
        except Exception as e:
            st.error(f"Couldn't read the file: {e}")
            jd_text = ""

        if not jd_text:
            st.warning("I couldn't extract any text. If it's a scanned PDF, try exporting a text-based PDF or DOCX.")
        else:
            st.subheader("📜 Extracted Job Description")
            st.text_area("Preview", jd_text, height=220, key="single_preview")

            if st.button("✨ Generate Advert", key="single_generate"):
                if not OPENAI_API_KEY:
                    st.error("OPENAI_API_KEY is not set. Add it in Manage app → Settings → Secrets.")
                else:
                    with st.spinner("Generating Neogen-style advert..."):
                        try:
                            advert_raw = generate_neogen_advert(jd_text)
                            advert = strip_markdown(advert_raw)
                        except Exception as e:
                            st.error(f"Generation failed: {e}")
                            advert = ""

                    if advert:
                        st.subheader("✅ Generated Job Advert (plain text)")
                        st.code(advert)  # copy button

                        edited = st.text_area("Edit before download (optional)", advert, height=260, key="single_edit")
                        use_text = edited if edited.strip() else advert

                        out_bytes = to_docx_structured(use_text)
                        st.download_button(
                            "⬇️ Download Advert as DOCX",
                            data=out_bytes,
                            file_name="neogen_job_advert.docx",
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                            key="single_download",
                        )

with tab_batch:
    files = st.file_uploader(
        "Upload multiple Job Descriptions (.docx or .pdf)",
        type=["docx", "pdf"],
        accept_multiple_files=True,
        key="batch_uploader",
    )
    st.caption("Tip: drag several files at once. You’ll get a ZIP of DOCX adverts.")

    if files and st.button("✨ Generate Adverts (Batch)", key="batch_generate"):
        if not OPENAI_API_KEY:
            st.error("OPENAI_API_KEY is not set. Add it in Manage app → Settings → Secrets.")
        else:
            zip_buf = BytesIO()
            try:
                with ZipFile(zip_buf, "w") as z:
                    for f in files:
                        try:
                            jd_text = extract_text_from_docx(f) if f.name.lower().endswith(".docx") \
                                      else extract_text_from_pdf(f)
                            if not jd_text:
                                continue
                            advert_raw = generate_neogen_advert(jd_text)
                            advert = strip_markdown(advert_raw)
                            docx_bytes = to_docx_structured(advert).getvalue()
                            base = os.path.splitext(os.path.basename(f.name))[0]
                            z.writestr(f"{base}_neogen_advert.docx", docx_bytes)
                        except Exception as e:
                            base = os.path.splitext(os.path.basename(f.name))[0]
                            z.writestr(f"{base}_ERROR.txt", f"Failed to process {f.name}:\n{e}")
                zip_buf.seek(0)
                st.success("Batch complete.")
                st.download_button(
                    "⬇️ Download ZIP of adverts",
                    data=zip_buf,
                    file_name="neogen_job_adverts.zip",
                    mime="application/zip",
                    key="batch_download",
                )
            except Exception as e:
                st.error(f"Batch failed: {e}")
